"""
generate_v541_final_theorem_bundle.py

Recreates the V541 final theorem bundle:
- V541_FINAL_THEOREM_REPORT.docx
- V541_FORMAL_THEOREM_REPORT.md
- V541_CLAIM_BOUNDARIES.md
- V541_VALIDATION_INDEX.md
- V541_conformal_recoverability_flow_infographic.png
- v541_final_theorem_bundle.zip

Expected prior archives in /mnt/data, if available:
- v513_eta_proof_package.zip
- v536_omega_convergence_outputs.zip
- v537_mu_defect_measure_outputs.zip
- v538_lyapunov_outputs.zip
"""

from pathlib import Path
import zipfile
import shutil
import textwrap
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

base = Path("/mnt/data")
pkg = base / "v541_final_theorem_bundle"
pkg.mkdir(exist_ok=True)

# -----------------------------
# Infographic
# -----------------------------
info_path = pkg / "V541_conformal_recoverability_flow_infographic.png"

img = Image.new("RGB", (1600, 1000), (246, 248, 251))
draw = ImageDraw.Draw(img)

draw.rectangle([40, 35, 1560, 135], fill=(24, 52, 91))
draw.text((70, 62), "V541 Final Theorem Bundle", fill="white")
draw.text((70, 98), "Conformal Recoverability Flow — Main theorem consolidation", fill=(220, 230, 242))

boxes = [
    ("Core Geometry", "g_eff(x,t) = Ω(x,t)^2 g₀(x)\nRecoverability becomes an effective conformal geometry."),
    ("Reserve Law", "C_t = M_tR_tL_t + λ₀η_convertB_t\nNominal branch volume counts only when convertible."),
    ("η_convert Closure", "η = harmonic serial conductance × exp[-minimum repair action]\nCloses the liquidity-trap gap."),
    ("Weak Evolution", "∂Ω/∂t = Source − Repair − μ_defect\nDefects enter as localized measure leakage."),
    ("Lyapunov Stability", "V = E[Ω] + reserve penalty + defect penalty + bottleneck penalty\nStable recovery requires constrained descent."),
    ("Evidence Chain", "V513 η proof | V536 Ω convergence | V537 μ measure convergence | V538 Lyapunov audit"),
]
coords = [(70,180),(835,180),(70,380),(835,380),(70,580),(835,580)]

for (title, body), (x,y) in zip(boxes, coords):
    draw.rounded_rectangle([x,y,x+690,y+150], radius=18, fill="white", outline=(120,145,175), width=2)
    draw.text((x+25,y+18), title, fill=(15, 40, 75))
    draw.text((x+25,y+55), textwrap.fill(body, width=62), fill=(35,35,35))

draw.rectangle([40, 825, 1560, 940], fill=(238, 243, 250), outline=(145,165,190), width=2)
draw.text((70, 850), "Theorem-shaped claim:", fill=(20,40,70))
draw.text(
    (70, 885),
    "Adaptive branch systems recover when conformal geometry relaxes while reserve remains above floor and defect-measure leakage does not grow.",
    fill=(30,30,30),
)
draw.text((70, 925), "Status: strong formal candidate; full uniqueness/rigorous continuum proof still open.", fill=(90,70,30))
img.save(info_path)

# -----------------------------
# Markdown theorem report
# -----------------------------
md = r"""# V541 Final Theorem Bundle

## Conformal Recoverability Flow

This package consolidates the retained-geometry theorem candidate after the η_convert closure, Ω convergence audit, μ_defect measure convergence audit, and Lyapunov stability audit.

## Core object

\[
g_{\mathrm{eff}}(x,t)=\Omega(x,t)^2g_0(x)
\]

The retained bridge appears to generate a conformal recoverability geometry. The scalar field \(\Omega\) controls effective distance, strain, repair paths, and curvature-like response.

## Reserve law

\[
C_t = M_tR_tL_t + \lambda_0\eta_{\mathrm{convert}}(t)B_t
\]

The term \(B_t\) is nominal recoverable branch volume. It contributes to usable reserve only through \(\eta_{\mathrm{convert}}\). This prevents the liquidity trap: large future volume does not imply actual safety unless the futures are convertible.

## η_convert closure

\[
\eta_{\mathrm{convert}}
=
\frac{\sum_i w_i}{\sum_i w_i/\eta_i}
\cdot
\exp[-C_{\mathrm{repair,min}}]
\]

where:

\[
w_i =
\frac{|\partial J/\partial \eta_i|}
{\sum_j|\partial J/\partial \eta_j|}
\]

and:

\[
C_{\mathrm{repair,min}}=\inf_u A_{\mathrm{repair}}[u]
\]

Interpretation:

\[
\eta_{\mathrm{convert}}
=
\text{effective serial recoverability conductance}
\times
\text{minimum-action repair survival}
\]

Frozen minimal channel basis:

1. conductance,
2. lineage continuity,
3. topology redundancy,
4. repair convertibility,
5. defect containment.

## Weak-form evolution

\[
\partial_t\Omega
=
\mathrm{Source}
-
\mathrm{Repair}
-
\mu_{\mathrm{defect}}
\]

with:

\[
\mathrm{Source}
=
G_L*
\left[
\frac{T_{\mathrm{retained}}}
{C_t-C_{\mathrm{floor}}+\epsilon}
\right]
\]

The more rigorous weak form is:

\[
\int \phi\,\partial_t\Omega\,dx
=
\int \phi\,\mathrm{Source}\,dx
-
\int \phi\,\mathrm{Repair}\,dx
-
\int \phi\,d\mu_{\mathrm{defect}}
\]

for smooth test functions \(\phi\).

## Defect measure

\(\mu_{\mathrm{defect}}\) behaves like localized measure leakage, not a smooth residual. V537 showed bounded defect mass, stable centroid, sharpening peak, and weak-form residual collapse when the measure term is included.

## Lyapunov candidate

\[
V[\Omega,C,\mu]
=
E[\Omega]
+
\alpha\max(0,C_{\mathrm{floor}}-C)^2
+
\beta\mu_{\mathrm{defect}}(\mathcal{X})
+
\gamma L_{\mathrm{bottleneck}}
\]

Stable recovery requires:

1. \(V\) decreases,
2. \(C_t>C_{\mathrm{floor}}\),
3. \(\mu_{\mathrm{defect}}\) is non-increasing,
4. bottleneck leakage remains bounded.

V538 showed energy alone was insufficient, while constrained \(V\) separated stable recovery from false recovery/collapse.

## Theorem-shaped statement

Let an adaptive branch system admit:

1. a baseline branch metric \(g_0\),
2. an effective conformal metric \(g_{\mathrm{eff}}=\Omega^2g_0\),
3. a recoverability reserve \(C_t\),
4. a weak-form evolution for \(\Omega\),
5. localized defect-measure leakage,
6. bounded bottleneck leakage,
7. a calibrated and identifiable \(\eta_{\mathrm{convert}}\).

If \(V[\Omega,C,\mu]\) is non-increasing, \(C_t\) remains above floor, \(\mu_{\mathrm{defect}}\) is non-increasing, and bottleneck leakage is bounded, then the trajectory remains in or enters the retained recovery basin.

## Evidence chain

- **V513 η proof:** η_convert operational/theorem-shaped closure.
- **V536 Ω convergence:** smooth bulk Ω converges under refinement.
- **V537 μ_defect convergence:** defect term behaves like localized measure.
- **V538 Lyapunov audit:** constrained V distinguishes stable recovery from false recovery/collapse.
- **V539 flow classification:** closest known family is constrained conformal gradient flow with measure-valued defect forcing.

## Current status

Supported:

- conformal geometry form \(g_{\mathrm{eff}}=\Omega^2g_0\),
- source/reserve driver,
- localized defect measure,
- η_convert liquidity closure,
- constrained Lyapunov stability signal.

Still open:

1. full uniqueness proof of η channel basis,
2. rigorous convergence theorem for \(\Omega\),
3. rigorous measure convergence theorem for \(\mu_{\mathrm{defect}}\),
4. formal Lyapunov theorem under admissible assumptions,
5. exact relation to known geometric flows.

## One-line summary

The retained bridge appears to generate conformal recoverability flow: usable future geometry is governed by source/reserve loading, recoverability liquidity, localized defect-measure leakage, and constrained Lyapunov descent.
"""
(pkg / "V541_FORMAL_THEOREM_REPORT.md").write_text(md)

# -----------------------------
# Claim boundaries
# -----------------------------
claim = """# V541 Claim Boundaries

Allowed claim:
The retained-geometry toy produces a strong theorem-shaped candidate: conformal recoverability flow with weak-form defect-measure leakage and constrained Lyapunov stability.

Do not claim:
- General Relativity derivation.
- Physical spacetime equivalence.
- Cosmology.
- Universality across all adaptive systems.
- Completed mathematical theorem.

Strongest accurate statement:
The retained bridge appears to generate conformal recoverability geometry where usable future branch volume is limited by recoverability liquidity, repair burden, defect leakage, and reserve-floor constraints.
"""
(pkg / "V541_CLAIM_BOUNDARIES.md").write_text(claim)

# -----------------------------
# Validation index
# -----------------------------
index = """# V541 Validation Index

This bundle consolidates:

1. V513 η_convert proof package
   - validates η as recoverability liquidity, not branch capacity.

2. V536 Ω convergence outputs
   - validates stable smooth-bulk Ω behavior.

3. V537 μ_defect measure outputs
   - validates localized measure-like defect leakage.

4. V538 Lyapunov outputs
   - validates constrained Lyapunov stability signal.

Recommended review order:
1. Read V541_FORMAL_THEOREM_REPORT.md
2. Review V541_CLAIM_BOUNDARIES.md
3. Inspect the infographic
4. Run/review the prior proof zips
"""
(pkg / "V541_VALIDATION_INDEX.md").write_text(index)

# -----------------------------
# DOCX report
# -----------------------------
doc = Document()
doc.add_heading("V541 Final Theorem Report", level=1)
p = doc.add_paragraph("Conformal Recoverability Flow")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(str(info_path), width=Inches(6.5))

sections = [
    ("Core Object", "g_eff(x,t) = Ω(x,t)^2 g₀(x). The retained bridge appears to generate a conformal recoverability geometry."),
    ("Reserve Law", "C_t = M_tR_tL_t + λ₀η_convertB_t. Branch volume contributes to usable reserve only when it is convertible."),
    ("η_convert Closure", "η_convert = (Σw_i)/(Σw_i/η_i) × exp[-C_repair_min]. This is effective serial recoverability conductance discounted by minimum-action repair survival."),
    ("Weak Evolution", "∂Ω/∂t = Source - Repair - μ_defect, with Source = G_L * [T_retained / (C_t - C_floor + ε)]."),
    ("Defect Measure", "μ_defect is a localized measure-like leakage term. V537 showed bounded mass, stable centroid, sharpening peak, and weak-form residual improvement."),
    ("Lyapunov Candidate", "V = E[Ω] + reserve-floor penalty + defect-mass penalty + bottleneck-leakage penalty. V538 showed constrained V separates stable recovery from false recovery/collapse better than energy alone."),
    ("Theorem-Shaped Claim", "If V decreases, C remains above floor, μ_defect is non-increasing, and bottleneck leakage is bounded, then the trajectory remains in or enters the retained recovery basin."),
    ("Status", "Strong formal candidate. Remaining gaps: uniqueness of η basis, rigorous Ω convergence, rigorous μ_defect convergence, full Lyapunov proof, and exact relation to known geometric flows."),
]

for h, body in sections:
    doc.add_heading(h, level=2)
    doc.add_paragraph(body)

doc_path = pkg / "V541_FINAL_THEOREM_REPORT.docx"
doc.save(doc_path)

# -----------------------------
# Copy prior archives if present
# -----------------------------
prior_files = [
    base / "v513_eta_proof_package.zip",
    base / "v536_omega_convergence_outputs.zip",
    base / "v537_mu_defect_measure_outputs.zip",
    base / "v538_lyapunov_outputs.zip",
]
for f in prior_files:
    if f.exists():
        shutil.copy(f, pkg / f.name)

# -----------------------------
# README
# -----------------------------
readme = """# V541 Final Theorem Bundle

Files:
- V541_FINAL_THEOREM_REPORT.docx
- V541_FORMAL_THEOREM_REPORT.md
- V541_CLAIM_BOUNDARIES.md
- V541_VALIDATION_INDEX.md
- V541_conformal_recoverability_flow_infographic.png
- Prior validation archives if available:
  - v513_eta_proof_package.zip
  - v536_omega_convergence_outputs.zip
  - v537_mu_defect_measure_outputs.zip
  - v538_lyapunov_outputs.zip

Purpose:
Package the current main theorem candidate after closing the η_convert operational gap and strengthening Ω, μ_defect, and Lyapunov evidence.
"""
(pkg / "README.md").write_text(readme)

# -----------------------------
# ZIP
# -----------------------------
zip_path = base / "v541_final_theorem_bundle.zip"
if zip_path.exists():
    zip_path.unlink()

with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
    for p in pkg.rglob("*"):
        z.write(p, arcname=p.relative_to(pkg))

print("Created:", zip_path)
